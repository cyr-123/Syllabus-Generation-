import streamlit as st
import base64
from docx import Document
import markdown
from bs4 import BeautifulSoup
import io
import os

from openai import OpenAI  # ✅ GPT-5 import

# -------------------------------------------------
# OpenAI client configuration
# -------------------------------------------------
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

MODEL_ID = "gpt-5"

# -------------------------------------------------
# Prompt generation function
# -------------------------------------------------
def createprompt(course_title, reference_textbook):
    return f"""
Assume you are an expert teacher with deep knowledge of outcomes-based education.

Generate a syllabus for the course: '{course_title}' using the textbook: '{reference_textbook}'.

## Course Description:
A concise paragraph (3–5 sentences) that clearly describes the course '{course_title}'.

## Learning Objectives:
Based on '{course_title}' generate a bulleted list of key learning objectives that align with Bloom's Taxonomy.

## Course Objective & Outcome Alignment Taxonomy:
Create a comprehensive alignment table showing how Course Learning Objectives (CLOs) map to Program Learning Outcomes (PLOs).

For each CLO, provide:
- A unique identifier
- A brief learning objective
- The category
- The aligned PLO

Ensure each CLO is:
- Specific and measurable
- Aligned with appropriate PLO numbers
- Categorized under Knowledge, Cognitive Skills, or Interpersonal Skills & Responsibility

## Course Modules:
Based on '{course_title}' generate a set of comprehensive modules for this course. Each module should take 1–2 weeks to cover.

For each module, use the Outcomes-Based Education (OBE) framework and Bloom's levels.

Ensure that the Desired Learning Outcomes (DLOs) are aligned with different levels of Bloom's Taxonomy:
- Remembering
- Understandingstreamlit run app.py

- Applying
- Analyzing
- Evaluating
- Creating

For each module, output a matrix as a table with columns:
- Desired Learning Outcomes (DLO)
- Course Content/Subject Matter
- Textbooks/References
- Outcomes-Based Teaching & Learning (OBTL)
- Assessment of Learning Outcomes (ALO)
- Resource Material
- Time Table

## Format Requirements:
- Do not include extra commentary, explanation, or code—only the final syllabus content.
"""

# -------------------------------------------------
# Query GPT-5
# -------------------------------------------------
def query_gpt5(prompt):
    try:
        response = client.responses.create(
            model=MODEL_ID,
            input=prompt
        )
        return response.output_text
    except Exception as e:
        raise Exception(f"GPT-5 API Error: {e}")

# -------------------------------------------------
# Convert markdown to Word document
# -------------------------------------------------
def markdown_to_word(markdown_content, filename="syllabus.docx"):
    try:
        html = markdown.markdown(markdown_content, extensions=['tables', 'extra'])
        soup = BeautifulSoup(html, 'html.parser')

        doc = Document()

        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                if element.get_text().strip():
                    doc.add_paragraph(element.get_text().strip())
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['th', 'td']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid'
                    for row_idx, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for col_idx, cell in enumerate(cells):
                            table.cell(row_idx, col_idx).text = cell.get_text().strip()

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        st.error(f"Error converting to Word: {e}")
        return None

# -------------------------------------------------
# Markdown file download
# -------------------------------------------------
def download_markdown(content):
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:file/markdown;base64,{b64}" download="syllabus.md">📄 Download Markdown</a>'
    return href

# -------------------------------------------------
# Main Streamlit App
# -------------------------------------------------
def main():
    st.title("📚 OBE Syllabus Generator (GPT-5)")

    with st.expander("ℹ️ About"):
        st.write("""
        This app generates a comprehensive course syllabus using GPT-5,
        structured using the Outcomes-Based Education (OBE) framework and Bloom's Taxonomy.
        """)

    course_title = st.text_input("🎓 Course Title", placeholder="e.g., Information Systems")
    reference_textbook = st.text_input("📘 Reference Textbooks", placeholder="e.g., Building Information Systems Using ML & DL")

    if st.button("🚀 Generate Course Syllabus"):
        if not course_title.strip():
            st.warning("⚠️ Please enter a course title.")
            return
        
        if not reference_textbook.strip():
            st.warning("⚠️ Please enter reference textbooks.")
            return

        prompt = createprompt(course_title, reference_textbook)

        try:
            with st.spinner(f"Generating syllabus for: {course_title}..."):
                response_text = query_gpt5(prompt)

                full_syllabus = f"# 📘 Syllabus for {course_title}\n\n"
                full_syllabus += f"**Reference Textbooks:** {reference_textbook}\n\n"
                full_syllabus += response_text

                st.markdown(full_syllabus)
                st.success("✅ Syllabus generated successfully!")

                st.markdown(download_markdown(full_syllabus), unsafe_allow_html=True)

                word_doc = markdown_to_word(full_syllabus)
                if word_doc:
                    st.download_button(
                        label="📄 Download Word Document",
                        data=word_doc.getvalue(),
                        file_name=f"{course_title.replace(' ', '_')}_syllabus.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            st.error(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
